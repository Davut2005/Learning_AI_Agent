"""
Text extraction from uploaded documents.

Supported formats:
  .txt  / .md              — plain UTF-8 read
  .pdf                     — pypdf page-by-page text extraction
  .docx                    — python-docx paragraphs + table cells
  .jpg/.jpeg/.png/.webp    — Tesseract OCR via pytesseract + Pillow

Tesseract must be installed as a system package (pytesseract is only the Python wrapper):
  Local (Windows)  →  https://github.com/UB-Mannheim/tesseract/wiki  (add install dir to PATH)
  Local (macOS)    →  brew install tesseract
  Local (Linux)    →  apt-get install -y tesseract-ocr
  Render           →  in Build Command add:
                         apt-get install -y tesseract-ocr tesseract-ocr-eng
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}


def extract_text(file_path: Path) -> str:
    """
    Extract all readable text from *file_path*.

    Returns the cleaned text string, or raises:
      ValueError   — unreadable / empty / unsupported content
      RuntimeError — missing system/library dependency
    """
    suffix = file_path.suffix.lower()

    if suffix in (".txt", ".md"):
        return _extract_plaintext(file_path)
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)
    if suffix in IMAGE_EXTENSIONS:
        return _extract_image(file_path)

    raise ValueError(
        f"Unsupported file type '{suffix}'. "
        "Supported: .txt, .md, .pdf, .docx, .jpg, .jpeg, .png, .webp"
    )


# ── Plain text ────────────────────────────────────────────────────────────────

def _extract_plaintext(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        raise ValueError(f"File '{path.name}' is empty.")
    return text


# ── PDF ───────────────────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("pypdf is not installed. Run: uv add pypdf") from e

    try:
        reader = PdfReader(str(path))
    except Exception as e:
        raise ValueError(f"Could not open PDF '{path.name}': {e}") from e

    if reader.is_encrypted:
        raise ValueError(
            f"PDF '{path.name}' is password-protected and cannot be processed."
        )

    pages_text: list[str] = []
    for page_num, page in enumerate(reader.pages, start=1):
        try:
            page_text = (page.extract_text() or "").strip()
            if page_text:
                pages_text.append(page_text)
        except Exception as e:
            logger.warning("PDF page %d extraction error in '%s': %s", page_num, path.name, e)

    if not pages_text:
        raise ValueError(
            f"PDF '{path.name}' contains no extractable text. "
            "It may be a scanned image-only PDF — convert it to a searchable PDF first."
        )

    full_text = "\n\n".join(pages_text)
    logger.info("PDF '%s': extracted %d chars from %d pages", path.name, len(full_text), len(pages_text))
    return full_text


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _extract_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise RuntimeError("python-docx is not installed. Run: uv add python-docx") from e

    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        raise ValueError(f"Could not open DOCX '{path.name}': {e}") from e

    parts: list[str] = []

    # Paragraphs (headings, body text, list items)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables — every cell, joined with a pipe separator per row
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))

    if not parts:
        raise ValueError(f"DOCX '{path.name}' contains no extractable text.")

    full_text = "\n\n".join(parts)
    logger.info("DOCX '%s': extracted %d chars from %d text blocks", path.name, len(full_text), len(parts))
    return full_text


# ── Image OCR (Tesseract) ─────────────────────────────────────────────────────

def _extract_image(path: Path) -> str:
    """
    Run Tesseract OCR on the image via pytesseract + Pillow.
    Tesseract must be installed separately (see module docstring).
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError as e:
        raise RuntimeError(
            "pytesseract or Pillow is not installed. Run: uv add pytesseract Pillow"
        ) from e

    try:
        img = Image.open(path)
    except Exception as e:
        raise ValueError(f"Could not open image '{path.name}': {e}") from e

    # Convert to RGB so Tesseract always gets a consistent colour space
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    try:
        # lang="eng" can be extended e.g. "eng+deu+fra" for multilingual docs
        text = pytesseract.image_to_string(img, lang="eng", config="--psm 3")
    except pytesseract.TesseractNotFoundError as e:
        raise RuntimeError(
            "Tesseract is not installed or not in PATH. "
            "Windows: https://github.com/UB-Mannheim/tesseract/wiki  "
            "macOS: brew install tesseract  "
            "Linux/Render: apt-get install -y tesseract-ocr"
        ) from e
    except Exception as e:
        raise RuntimeError(f"Tesseract OCR failed for '{path.name}': {e}") from e

    text = text.strip()
    if not text:
        raise ValueError(
            f"Image '{path.name}' contains no readable text "
            "(blank image or unrecognised content)."
        )

    logger.info("Image OCR '%s': extracted %d chars via Tesseract", path.name, len(text))
    return text
